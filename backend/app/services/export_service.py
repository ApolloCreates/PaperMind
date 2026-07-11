from sqlalchemy.orm import Session

from app.repositories.draft_repository import DraftRepository
from app.repositories.draft_section_repository import (
    DraftSectionRepository,
)
from io import BytesIO

from docx import Document

from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
)

from reportlab.lib.styles import getSampleStyleSheet


class ExportService:

    def __init__(self):

        self.draft_repo = DraftRepository()
        self.section_repo = DraftSectionRepository()

    def build_markdown(
        self,
        db: Session,
        draft_id: str,
    ) -> str:

        draft = self.draft_repo.get(
            db,
            draft_id,
        )

        if draft is None:
            raise ValueError("Draft not found.")

        sections = self.section_repo.list_sections(
            db,
            draft_id,
        )

        if not sections:
            raise ValueError("Draft has no sections.")

        md = f"# {draft.title}\n\n"

        for section in sections:

            md += (
                f"## {section.section.value}\n\n"
                f"{section.content}\n\n"
            )

        return md
    
    
    def build_docx(
        self,
        db: Session,
        draft_id: str,
    ) -> BytesIO:

        draft = self.draft_repo.get(
            db,
            draft_id,
        )

        if draft is None:
            raise ValueError("Draft not found.")

        sections = self.section_repo.list_sections(
            db,
            draft_id,
        )

        if not sections:
            raise ValueError("Draft has no sections.")

        document = Document()

        document.add_heading(
            draft.title,
            level=0,
        )

        for section in sections:

            document.add_heading(
                section.section.value,
                level=1,
            )

            document.add_paragraph(
                section.content,
            )

        buffer = BytesIO()

        document.save(buffer)

        buffer.seek(0)

        return buffer
    
    
    
    def build_pdf(
        self,
        db: Session,
        draft_id: str,
    ) -> BytesIO:

        draft = self.draft_repo.get(
            db,
            draft_id,
        )

        if draft is None:
            raise ValueError("Draft not found.")

        sections = self.section_repo.list_sections(
            db,
            draft_id,
        )

        if not sections:
            raise ValueError("Draft has no sections.")

        buffer = BytesIO()

        styles = getSampleStyleSheet()

        doc = SimpleDocTemplate(buffer)

        story = []

        story.append(
            Paragraph(
                draft.title,
                styles["Title"],
            )
        )

        for section in sections:

            story.append(
                Paragraph(
                    section.section.value,
                    styles["Heading1"],
                )
            )

            story.append(
                Paragraph(
                    section.content.replace(
                        "\n",
                        "<br/>",
                    ),
                    styles["BodyText"],
                )
            )

        doc.build(story)

        buffer.seek(0)

        return buffer